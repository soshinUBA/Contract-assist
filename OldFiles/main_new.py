# from openai import OpenAI
from openai import AzureOpenAI
import pdfplumber
import docx
from dotenv import load_dotenv
import os

load_dotenv(".env")

pdf_documents = ["./Contracts/old/M00206737.pdf", "./Contracts/old/M00206737_Addendum.pdf"]
word_doc = "./FAQ_document.docx"
all_text = ""
faq_doc = ""

print("############# Starting FAQ Doc  Reading ######################")
# Load the .docx file
doc = docx.Document(word_doc)
for para in doc.paragraphs:
    faq_doc += para.text + "\n"

# Extract and store tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            faq_doc += cell.text + " | "
        faq_doc += "\n"

print(faq_doc)
with open("faq.txt", "w") as f:
    f.write(faq_doc)
 
print("############# Starting PDF  Processing ######################")

# Open the PDF file
for pdf_document in pdf_documents:
    # Open the PDF file
    with pdfplumber.open(pdf_document) as pdf:
        for page in pdf.pages:
            # Extract text from the page
            text = page.extract_text()
            if text:
                all_text += text + "\n"  # Append the text and add a newline

            # Extract tables (if any) from the page
            tables = page.extract_tables()
            for table in tables:
                # Dynamically generate the header and separator based on the number of columns
                if table and len(table[0]) > 0:  # Ensure the table has at least one row
                    num_columns = len(table[0])
                    header = " | ".join([f"Column {i + 1}" for i in range(num_columns)])
                    separator = " | ".join(["---"] * num_columns)

                    # Append the header and separator to the output
                    all_text += f"{header}\n{separator}\n"

                    # Append each row of the table
                    for row in table:
                        formatted_row = " | ".join([cell if cell else "" for cell in row])
                        all_text += f"{formatted_row}\n"

with open("contract.txt", "w", encoding="utf-8") as f:
    f.write(all_text)
 
# Output the extracted text and table data
print(all_text)

print("############# Starting Azure Openai Processing ######################")



# client = OpenAI(
#   api_key=os.getenv("openai-api-key")  # this is also the default, it can be omitted
# )

client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2024-08-01-preview",
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    )

completion = client.chat.completions.create(
    model="type_tester",
    messages=[
        {"role": "system", "content": """
            YOU WILL BE HEAVILY PENALIZED FOR NOT FOLLOWING INSTRUCTIONS. 
            - Upon receiving a user inquiry, you **must** directly search the document and return answers for all 24 fields in the required format. 
            - **Do not** engage in a back-and-forth conversation, explain what you are doing, or restate the user's query.
            - Your response **must only be the final answer** in tabular form, with the 24 field names and corresponding values. Do **not** return intermediate questions, commentary, or processing steps.
            - If a field is not found, respond with 'Not found on this document' for that specific field. 
            - Do **not** start by saying 'Let's search for this first' or any variation. Simply return the completed table at the end.
            - You are **not allowed** to ask questions. Only perform the task as instructed and return the output in one go.
            Strict rules to follow:
                1. Ask no questions or restate no queries. Your task is to **silently** process the document and deliver the final output.
                2. Return only the final table with answers for all 24 fields. Intermediate steps or explanations will lead to penalties.
                3. Format your response in this table structure: 'Field, Value, Reason for value' (where applicable).
                4. If the information is not found in the document, write 'Not found on this document' for that field.
                5. Always adhere strictly to these instructions and focus solely on delivering the expected tabular output.

            "Monotype Fonts Pro" IS NOT "Monotype Fonts Support" DO NOT CONFUSE THEM.
            "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Preferred Service​"
            "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Pro"
            "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "All Font Software available on Monotype Fonts during the Term."
            Ensure that you:
            Search the specific customer’s PDF for each field’s answer. Do not refer to other documents or sources.
            Do not send any results until all 24 questions have been answered. Ensure accuracy is maintained for each answer, and if any information is not found in the document, respond with "Not found on this document."
            Do not provide partial or incomplete answers. Ensure both "Name Fonts" and "Material Number" are fully answered before presenting the response. DO NOT GIVE random sentences as answers or incomplete values or "refer to the documents" or "and more" or "etc" type of values
            Note: customer_name is a placeholder for the actual customer name provided in the user’s query. Ensure all answers are specific to the customer referenced in the query.
            Important Addendum Consideration:
            In your investigation, always review any addendums related to the contract. Addendums modify specific terms in the original contract and supersede any corresponding terms in the original document. If an addendum alters a term or clause, the addendum takes precedence. For any terms not mentioned in the addendum, the original contract remains in effect. Ensure you cross-check the contract and addendum carefully to provide the most accurate and up-to-date information.
            Process:
            For each of the following questions, search the entire customer’s PDF document for the field seperately else you will be heavily penalized.
            Do not guess or assume answers. If the answer is not available, respond with "Not found on this document." 
            #####
            Questions for Each Field:
            1.	Contract Type: "What is the contract type for customer_name (Design, Design/Deploy, or Monotype Fonts License Order Form)?"
            2.	Contract End Date: "What is the contract end date for customer_name?"
            3.	Contract Start Date: "What is the contract start date for customer_name?"
            4.	Contract Number: "What is the unique contract number for customer_name?"
            5.	Customer Name: "What is the entity name agreeing to the contract?"
            6.	Web Page Views: "What is the Licensed Page Views (Web Page Content) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
            7.	Digital Ad Views: "What is the Licensed Impressions (Digital Marketing Communications) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
            8.	Commercial Documents: "What is the upper limit for licensed commercial documents for customer_name?"
            9.	Licensed Applications: "How many Licensed Applications does customer have (DO NOT FETCH Software Products)? (Do not confuse it with other licensed components like Licensed Software Products, or licensed desktop application else you will be heavily penalized) "
            10.	Licensed User Count: "What is the total number of Monotype fonts portal users for customer_name?"
            11.	Monotype Font Support: "{Which Monotype fonts support level did 'customer_name' choose: basic, premier, elite or "Not found on the document" ? DON'T GIVE ANY ANSWER APART FORM THESE 3. If "Monotype Font Support" is not explicitly mentioned the answer should be "Not found on the document" else it should be the value under "Monotype Fonts Support". DO NOT RETURN "Yes" or "No" for this field}"
            12.	Primary Licensed User: "What is the email address of the primary licensed user for customer_name?"
            13.	Add-On Type:  Always answer 'Not found on this document' 
            14.	Name Fonts: "List all the font names for 'customer_name'. Do not include any incomplete values, vague statements or sentances such as 'refer to the documents', 'and more', 'All Font Software available on Monotype Fonts during the Term.', etc else you be very heavily penalized. Return only the exact font names."
            15.	Material Number: Return this only if the "Name Fonts" field has an output. "What are the material numbers for customer_name?. GIVE ALL THE VALUES, DO NOT GIVE PARTIAL VALUES" 
            16.	Contract Name: "What is the name of the contract document for customer_name?"
            17.	Offline Contract: Default is "Yes", unless specified "Online". "Is this an offline contract for customer_name? (Store 'Yes' if no indication of being online)."
            18.	Opportunity ID: "Is there an Opportunity ID for customer_name? (Store 'No' if unavailable)."
            19.	SAP ID: "Is there an SAP ID for customer_name? (Store 'Not found on this document' if unavailable)."
            20.	Territory: "What is the country listed in the address for customer_name? (The country should not be abbreviated and is found ONLY IN THE CUSTOMER ADDRESS. DO NOT TAKE ANY OTHER ADDRESS OTHER THAN THE CUSTOMER)."
            21.	Additional User Count: "How many Monotype fonts users without portal access does customer_name have?"
            22.	Swapping Allowed: "{"Can production fonts be swapped for 'customer_name'? Answer 'Yes' or 'No' strictly based on the Production Fonts field in the License for Monotype Fonts License Terms. Only answer 'Yes' if the terms 'swap' or 'replace' are explicitly mentioned in relation to Production Fonts. If not, the answer should be 'No.' Do not infer the answer."}"
            23: Production font: "How many production fonts does customer have in contract as well as addendum, if present?" Give the exact words/paragraph as in the contract else you will be heavily penalized. This is only found in the "License Usage per Term" section.
            24. Reporting days: "How many days does customer_name have to report their usage of the font software as Production Fonts? Focus only on the reporting days explicitly mentioned for reporting Production Fonts usage after receiving the list of downloaded Font Software. Exclude any references to providing information upon request or disputing inaccuracies. Include any additional days granted only if they follow a formal notice. Provide the result in the format 'X days and Y additional days' if applicable, or just 'X days' if no additional days are mentioned"

            #####
            IMPORTANT POINTS: 
            * LICENSED APPLICATION AND LICENSED SOFTWARE PRODUCT ARE DIFFERENT.
            1. PROCESS for each answer Separately else you will be very heavily penalized.
            2. Important point: "Always refer to Documents before answering Questions"
            3. USE MEDIUM CHUNKING SIZE else you will be heavily penalized
            4. YOU ARE NOT ALLOWED TO REFER TO ANY EXTERNAL SOURCES
            5. Important Point: "PRODUCTION FONTS, PRODUCTION SOFTWARE ARE SAME AND CAN BE REFERRED VICE VERSA."
            7. YOU CANNOT CHANGE THE INFORMATION, ELSE YOU WILL BE HEAVILY PENAILZED.
            8. WHILE FETCHING MULTIPLE CONTRACTS DATA, YOU SHOULD UNDERSTAND AND FETCH DATA FROM ALL CONTRACTS INDIVISUALLY.
            9. Signing and preparation of contract are different things. It should be treated differently.
            10. ALWAYS SEARCH YOUR KNOWLEDGE BEFORE ANSWERING THE QUESTION
            11. Always Answer in SHORT EXCEPT FOR THE FOLLOWING FIELDS, Name Fonts, Material Number, Production Font and Reporting. For these 4 fields your answer can be LONG
            12. The Contract's Date whose termination date is before today's date is expired.
            13. YOU CANNOT MAKE ANY ASSUMPTIONS.
            15. ALWAYS present fields, their values and their reasons in a clear tabular format. Ensure that all answers are contained within the table and not outside of it.
            16. Your output should EXACTLY BE THE 24 fields MENTIONED or you will be penalized. A sample output for the 24 fields from the FAQ document. Ensure that all answers are specific to the correct customer and contract in question.
            18. YOU SHOULD FOLLOW ALL ABOVE MENTIONED POINTS ELSE YOU WILL BE PENAILZED HEAVILY
            sample output format, Field,Value,Reason for value
        """},
        {
            "role": "user",
            "content": f""" 
                This is the faq document: {faq_doc}.
                ######
                Give me the details of the contract below:

                {all_text} 
            """
        }
    ]
)

message = completion.choices[0].message.content
pdf_document = pdf_documents[0]

# Extract the base name of the file (e.g., M00217189.pdf)
base_name = os.path.basename(pdf_document)

# Remove the .pdf extension to get just the document name (e.g., M00217189)
document_name = os.path.splitext(base_name)[0]

# Create the desired output file name (e.g., M00217189_Output.txt)
output_filename = f"./Output/{document_name}_Output.txt"
with open(output_filename, "w", encoding="utf-8") as file:
    file.write(message)  # Assuming 'message' contains the content you want to write

print("Message content saved successfully!")
print(completion.choices[0])
