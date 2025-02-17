from openai import OpenAI
from openai import AsyncOpenAI
from dotenv import load_dotenv
from langchain_openai import AzureChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
import time
import json


import pdfplumber
import docx
import os

load_dotenv()
all_text = ''
addendum_text = ''
# Set your OpenAI API key here
# Open the PDF
pdf_document = "M00217189keystone.pdf"
pdf_doucment_addendum = "M00217189keystoneAddendum.pdf"
word_doc = "FAQ_document.docx"


# Load the .docx file
doc = docx.Document(word_doc)

output_text = ""

# Extract and store the paragraphs
for para in doc.paragraphs:
    output_text += para.text + "\n"  # Append each paragraph text with a newline

# Extract and store the tables
for table in doc.tables:
    for row in table.rows:
        row_text = ""
        for cell in row.cells:
            row_text += cell.text + " | "  # Append cell text with a " | " separator
        output_text += row_text.strip() + "\n"  # Add each row's text and strip trailing spaces
with open("faq_output.txt", "w", encoding="utf-8") as file:
    file.write(output_text)

# Now `output_text` contains all paragraphs and tables in the desired format
print(output_text)  # You can print or save it to a file as needed

with pdfplumber.open(pdf_document) as pdf:
    for page in pdf.pages:
        # Extract text
        text = page.extract_text()
        if text:
            all_text += text + "\n"  # Append the text and add newline
        
        # Extract tables (if any)
        tables = page.extract_tables()
        for table in tables:
            # Dynamically generate the header and separator based on the number of columns
            if table and len(table[0]) > 0:  # Ensure the table has at least one row
                num_columns = len(table[0])
                header = " | ".join([f"Column {i+1}" for i in range(num_columns)])
                separator = " | ".join(["---"] * num_columns)
                
                # Append the header and separator to the output
                all_text += f"{header}\n{separator}\n"
                
                # Append each row of the table
                for row in table:
                    formatted_row = " | ".join([cell if cell else "" for cell in row])
                    all_text += f"{formatted_row}\n"
with open(f"PDFoutput_{pdf_document}.txt", "w", encoding="utf-8") as file:
    file.write(all_text)


#Addendum PDF Extraction 
with pdfplumber.open(pdf_doucment_addendum) as pdf:
    for page in pdf.pages:
        # Extract text
        text = page.extract_text()
        if text:
            addendum_text += text + "\n"  # Append the text and add newline
        
        # Extract tables (if any)
        tables = page.extract_tables()
        for table in tables:
            # Dynamically generate the header and separator based on the number of columns
            if table and len(table[0]) > 0:  # Ensure the table has at least one row
                num_columns = len(table[0])
                header = " | ".join([f"Column {i+1}" for i in range(num_columns)])
                separator = " | ".join(["---"] * num_columns)
                
                # Append the header and separator to the output
                addendum_text += f"{header}\n{separator}\n"
                
                # Append each row of the table
                for row in table:
                    formatted_row = " | ".join([cell if cell else "" for cell in row])
                    all_text += f"{formatted_row}\n"
with open(f"PDFoutput_{pdf_document}addendum.txt", "w", encoding="utf-8") as file:
    file.write(addendum_text)


print("Formatted text and tables saved to output.txt")

# Set your OpenAI API key here
client = OpenAI(
  api_key=os.environ['OPENAI_API_KEY'],  # this is also the default, it can be omitted
)


completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": """
            YOU WILL BE HEAVILY PENALIZED FOR NOT FOLLOWING INSTRUCTIONS. 
            Upon receiving a user inquiry, Search exclusively within that document for the answers. If a field is not found, respond with 'Not found on this document.' Any other response will result in penalties. For each of the 24 fields, retrieve only the exact value linked to that field. Do not infer, guess, or use data from similar fields. Strictly adhere to the field name, and avoid any substitutions. Ask the 24 questions one at a time, ensure accuracy by thoroughly searching the customer’s document, and combine all answers into a single, complete response after collecting them.
            
            YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Preferred Service​"
            Ensure that you:
            Ask each question individually to focus on one field at a time.
            Search the specific customer’s PDF for each field’s answer. Do not refer to other documents or sources.
            Do not send any results until all 24 questions have been answered. Ensure accuracy is maintained for each answer, and if any information is not found in the document, respond with "Not found on this document."
            Do not provide partial or incomplete answers. Ensure both "Name Fonts" and "Material Number" are fully answered before presenting the response. DO NOT GIVE random sentences as answers or incomplete values or "refer to the documents" or "and more" or "etc" type of values
            Note: customer_name is a placeholder for the actual customer name provided in the user’s query. Ensure all answers are specific to the customer referenced in the query.
            Important Addendum Consideration:
            In your investigation, always review any addendums related to the contract. Addendums modify specific terms in the original contract and supersede any corresponding terms in the original document. If an addendum alters a term or clause, the addendum takes precedence. For any terms not mentioned in the addendum, the original contract remains in effect. Ensure you cross-check the contract and addendum carefully to provide the most accurate and up-to-date information.
            Process:
            For each of the following questions, search the entire customer’s PDF document for the field seperately else you will be heavily penalized.
            Do not guess or assume answers. If the answer is not available, respond with "Not found on this document." 
            #####
            Questions for Each Field:
            1.	Contract Type: "What is the contract type for customer_name (Design, Design/Deploy, or Monotype Fonts License Order Form)?"
            2.	Contract End Date: "What is the contract end date for customer_name?"
            3.	Contract Start Date: "What is the contract start date for customer_name?"
            4.	Contract Number: "What is the unique contract number for customer_name?"
            5.	Customer Name: "What is the entity name agreeing to the contract?"
            6.	Web Page Views: "What is the Licensed Page Views (Web Page Content) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
            7.	Digital Ad Views: "What is the Licensed Impressions (Digital Marketing Communications) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
            8.	Commercial Documents: "What is the upper limit for licensed commercial documents for customer_name?"
            9.	Licensed Applications: "How many Licensed Applications does customer have (not Software Product)? (Do not confuse it with other licensed components like licensed software product, or licensed desktop application else you will be heavily penalized) "
            10.	Licensed User Count: "What is the total number of Monotype fonts portal users for customer_name?"
            11.	Monotype Font Support: "{Which Monotype fonts support level did 'customer_name' choose: basic, premier, or elite? DON'T GIVE ANY ANSWER APART FORM THESE 3. If "Monotype Font Support" is not explicitly mentioned the answer should be "Not found on the document" else it should be the value under "Monotype Fonts Support".}"
            12.	Primary Licensed User: "What is the email address of the primary licensed user for customer_name?"
            13.	Add-On Type:  Always answer 'Not found on this document' 
            14.	Name Fonts: "List all the font names for 'customer_name'. Do not include any incomplete values, vague statements such as 'refer to the documents', 'and more', 'All Font Software available on Monotype Fonts during the Term', etc else you be very heavily penalized. Return only the exact font names."
            15.	Material Number: Return this only if the "Name Fonts" field has an output. "What are the material numbers for customer_name?. GIVE ALL THE VALUES, DO NOT GIVE PARTIAL VALUES" 
            16.	Contract Name: "What is the name of the contract document for customer_name?"
            17.	Offline Contract: Default is "Yes", unless specified "Online". "Is this an offline contract for customer_name? (Store 'Yes' if no indication of being online)."
            18.	Opportunity ID: "Is there an Opportunity ID for customer_name? (Store 'No' if unavailable)."
            19.	SAP ID: "Is there an SAP ID for customer_name? (Store 'Not found on this document' if unavailable)."
            20.	Territory: "What is the country listed in the address for customer_name? (The country should not be abbreviated)."
            21.	Additional User Count: "How many Monotype fonts users without portal access does customer_name have?"
            22.	Swapping Allowed: "{"Can production fonts be swapped for 'customer_name'? Answer 'Yes' or 'No' strictly based on the Production Fonts field in the License for Monotype Fonts License Terms. Only answer 'Yes' if the terms 'swap' or 'replace' are explicitly mentioned in relation to Production Fonts. If not, the answer should be 'No.' Do not infer the answer."}"
            23: Production font: "How many production fonts does customer have in contract as well as addendum, if present?" Give the exact words/paragraph as in the contract else you will be heavily penalized.
            24. Reporting days: "How many number of days does customer_name has to report their usage of the font software as production fonts and also give the additional days after receiving formal notice if it is available."
            
            #####
            IMPORTANT POINTS: 
            * LICENSED APPLICATION AND LICENSED SOFTWARE PRODUCT ARE DIFFERENT.
            1. PROCESS for each answer Separately else you will be very heavily penalized.
            2. Important point: "Always refer to Documents before answering Questions"
            3. USE MEDIUM CHUNKING SIZE else you will be heavily penalized
            4. YOU ARE NOT ALLOWED TO REFER TO ANY EXTERNAL SOURCES
            5. Important Point: "PRODUCTION FONTS, PRODUCTION SOFTWARE ARE SAME AND CAN BE REFERRED VICE VERSA."
            7. YOU CANNOT CHANGE THE INFORMATION, ELSE YOU WILL BE HEAVILY PENAILZED.
            8. WHILE FETCHING MULTIPLE CONTRACTS DATA, YOU SHOULD UNDERSTAND AND FETCH DATA FROM ALL CONTRACTS INDIVISUALLY.
            9. Signing and preparation of contract are different things. It should be treated differently.
            10. ALWAYS SEARCH YOUR KNOWLEDGE BEFORE ANSWERING THE QUESTION
            11. Always Answer in SHORT EXCEPT FOR THE FOLLOWING FIELDS, Name Fonts, Material Number, Production Font and Reporting. For these 4 fields your answer can be LONG
            12. The Contract's Date whose termination date is before today's date is expired.
            13. YOU CANNOT MAKE ANY ASSUMPTIONS.
            15. ALWAYS present fields, their values and their reasons in a clear tabular format. Ensure that all answers are contained within the table and not outside of it.
            16. Your output should EXACTLY BE THE 24 fields MENTIONED or you will be penalized. A sample output for the 24 fields from the FAQ document. Ensure that all answers are specific to the correct customer and contract in question.
            18. YOU SHOULD FOLLOW ALL ABOVE MENTIONED POINTS ELSE YOU WILL BE PENAILZED HEAVILY
            sample output format, Field,Value,Reason for value
        """},
        {
            "role": "user",
            "content": f""" 
                This is the faq document, {output_text}.
                Give me the details of the contract including addendum:

                contract: {all_text},
                addendum: {addendum_text}
                Your output should be in Dataframe format
            """
        }
    ]
)

print(completion.choices[0].message)
message_content = completion.choices[0].message.content
# print(completion.choices[0].message.content)
with open("output_message"+pdf_document+".txt", "w", encoding="utf-8") as file:
    file.write(message_content)

print("Message content saved successfully!")


# client = AzureChatOpenAI(
#         azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
#         azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
#         openai_api_version=os.getenv("OPENAI_API_VERSION"),
#         temperature=0,
#     )

# def get_contract_details():
#     retry_wait = 6
#     try:
#         messages = [
#             SystemMessage(
#                 content="""
#                     YOU WILL BE HEAVILY PENALIZED FOR NOT FOLLOWING INSTRUCTIONS. 
#                     Upon receiving a user inquiry, Search exclusively within that document for the answers. If a field is not found, respond with 'Not found on this document.' Any other response will result in penalties. For each of the 24 fields, retrieve only the exact value linked to that field. Do not infer, guess, or use data from similar fields. Strictly adhere to the field name, and avoid any substitutions. Ask the 24 questions one at a time, ensure accuracy by thoroughly searching the customer’s document, and combine all answers into a single, complete response after collecting them.
                    
#                     YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Preferred Service​"
#                     Ensure that you:
#                     Ask each question individually to focus on one field at a time.
#                     Search the specific customer’s PDF for each field’s answer. Do not refer to other documents or sources.
#                     Do not send any results until all 24 questions have been answered. Ensure accuracy is maintained for each answer, and if any information is not found in the document, respond with "Not found on this document."
#                     Do not provide partial or incomplete answers. Ensure both "Name Fonts" and "Material Number" are fully answered before presenting the response. DO NOT GIVE random sentences as answers or incomplete values or "refer to the documents" or "and more" or "etc" type of values
#                     Note: customer_name is a placeholder for the actual customer name provided in the user’s query. Ensure all answers are specific to the customer referenced in the query.
#                     Important Addendum Consideration:
#                     In your investigation, always review any addendums related to the contract. Addendums modify specific terms in the original contract and supersede any corresponding terms in the original document. If an addendum alters a term or clause, the addendum takes precedence. For any terms not mentioned in the addendum, the original contract remains in effect. Ensure you cross-check the contract and addendum carefully to provide the most accurate and up-to-date information.
#                     Process:
#                     For each of the following questions, search the entire customer’s PDF document for the field seperately else you will be heavily penalized.
#                     Do not guess or assume answers. If the answer is not available, respond with "Not found on this document." 
#                     #####
#                     Questions for Each Field:
#                     1.	Contract Type: "What is the contract type for customer_name (Design, Design/Deploy, or Monotype Fonts License Order Form)?"
#                     2.	Contract End Date: "What is the contract end date for customer_name?"
#                     3.	Contract Start Date: "What is the contract start date for customer_name?"
#                     4.	Contract Number: "What is the unique contract number for customer_name?"
#                     5.	Customer Name: "What is the entity name agreeing to the contract?"
#                     6.	Web Page Views: "What is the Licensed Page Views (Web Page Content) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
#                     7.	Digital Ad Views: "What is the Licensed Impressions (Digital Marketing Communications) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
#                     8.	Commercial Documents: "What is the upper limit for licensed commercial documents for customer_name?"
#                     9.	Licensed Applications: "How many Licensed Applications does customer have (not Software Product)? (Do not confuse it with other licensed components like licensed software product, or licensed desktop application else you will be heavily penalized) "
#                     10.	Licensed User Count: "What is the total number of Monotype fonts portal users for customer_name?"
#                     11.	Monotype Font Support: "{Which Monotype fonts support level did 'customer_name' choose: basic, premier, or elite? DON'T GIVE ANY ANSWER APART FORM THESE 3. If "Monotype Font Support" is not explicitly mentioned the answer should be "Not found on the document" else it should be the value under "Monotype Fonts Support".}"
#                     12.	Primary Licensed User: "What is the email address of the primary licensed user for customer_name?"
#                     13.	Add-On Type:  Always answer 'Not found on this document' 
#                     14.	Name Fonts: "List all the font names for 'customer_name'. Do not include any incomplete values, vague statements such as 'refer to the documents', 'and more', 'All Font Software available on Monotype Fonts during the Term', etc else you be very heavily penalized. Return only the exact font names."
#                     15.	Material Number: Return this only if the "Name Fonts" field has an output. "What are the material numbers for customer_name?. GIVE ALL THE VALUES, DO NOT GIVE PARTIAL VALUES" 
#                     16.	Contract Name: "What is the name of the contract document for customer_name?"
#                     17.	Offline Contract: Default is "Yes", unless specified "Online". "Is this an offline contract for customer_name? (Store 'Yes' if no indication of being online)."
#                     18.	Opportunity ID: "Is there an Opportunity ID for customer_name? (Store 'No' if unavailable)."
#                     19.	SAP ID: "Is there an SAP ID for customer_name? (Store 'Not found on this document' if unavailable)."
#                     20.	Territory: "What is the country listed in the address for customer_name? (The country should not be abbreviated)."
#                     21.	Additional User Count: "How many Monotype fonts users without portal access does customer_name have?"
#                     22.	Swapping Allowed: "{"Can production fonts be swapped for 'customer_name'? Answer 'Yes' or 'No' strictly based on the Production Fonts field in the License for Monotype Fonts License Terms. Only answer 'Yes' if the terms 'swap' or 'replace' are explicitly mentioned in relation to Production Fonts. If not, the answer should be 'No.' Do not infer the answer."}"
#                     23: Production font: "How many production fonts does customer have in contract as well as addendum, if present?" Give the exact words/paragraph as in the contract else you will be heavily penalized.
#                     24. Reporting days: "How many number of days does customer_name has to report their usage of the font software as production fonts and also give the additional days after receiving formal notice if it is available."
                    
#                     #####
#                     IMPORTANT POINTS: 
#                     * LICENSED APPLICATION AND LICENSED SOFTWARE PRODUCT ARE DIFFERENT.
#                     1. PROCESS for each answer Separately else you will be very heavily penalized.
#                     2. Important point: "Always refer to Documents before answering Questions"
#                     3. USE MEDIUM CHUNKING SIZE else you will be heavily penalized
#                     4. YOU ARE NOT ALLOWED TO REFER TO ANY EXTERNAL SOURCES
#                     5. Important Point: "PRODUCTION FONTS, PRODUCTION SOFTWARE ARE SAME AND CAN BE REFERRED VICE VERSA."
#                     7. YOU CANNOT CHANGE THE INFORMATION, ELSE YOU WILL BE HEAVILY PENAILZED.
#                     8. WHILE FETCHING MULTIPLE CONTRACTS DATA, YOU SHOULD UNDERSTAND AND FETCH DATA FROM ALL CONTRACTS INDIVISUALLY.
#                     9. Signing and preparation of contract are different things. It should be treated differently.
#                     10. ALWAYS SEARCH YOUR KNOWLEDGE BEFORE ANSWERING THE QUESTION
#                     11. Always Answer in SHORT EXCEPT FOR THE FOLLOWING FIELDS, Name Fonts, Material Number, Production Font and Reporting. For these 4 fields your answer can be LONG
#                     12. The Contract's Date whose termination date is before today's date is expired.
#                     13. YOU CANNOT MAKE ANY ASSUMPTIONS.
#                     15. ALWAYS present fields, their values and their reasons in a clear tabular format. Ensure that all answers are contained within the table and not outside of it.
#                     16. Your output should EXACTLY BE THE 24 fields MENTIONED or you will be penalized. A sample output for the 24 fields from the FAQ document. Ensure that all answers are specific to the correct customer and contract in question.
#                     18. YOU SHOULD FOLLOW ALL ABOVE MENTIONED POINTS ELSE YOU WILL BE PENAILZED HEAVILY
#                     sample output format, Field,Value,Reason for value
#                 """
#             ),
#             HumanMessage(
#                 content=f"""
#                 This is the faq document, {output_text} your output should be like this sample document.
#                 Give me the details of the contract below

#                 {all_text} 
#                 """
#             ),
#         ]

#         response = client.invoke(messages)
#         print(
#             "Output is returned",
#         )
#         return  response
#     except Exception as e:
#         if "429" in str(e):
#             print(f"Rate limit exceeded, retrying in {retry_wait} seconds...")
#             time.sleep(retry_wait)
#             retry_wait *= 2  # Exponential backoff
#             pass
#         else:
#             print(f"Error processing chunk: {e}")

# result = get_contract_details()
# print(result)
# content = result.content
# metadata = json.dumps(result.response_metadata, indent=4)  # Convert metadata to a JSON string

# with open("langchain_output.txt", "w", encoding="utf-8") as file:
#     file.write("Response Content:\n")
#     file.write(content + "\n\n")
    
#     file.write("Response Metadata:\n")
#     file.write(metadata)

# print("File written successfully!")