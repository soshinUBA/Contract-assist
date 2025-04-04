from langchain_openai import AzureChatOpenAI
from dotenv import load_dotenv
import os
import tiktoken
import docx
import re
import boto3
from botocore.exceptions import ClientError
import string
load_dotenv()

def get_model():
    client = AzureChatOpenAI(
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
            openai_api_version=os.getenv("OPENAI_API_VERSION"),
            temperature=0,
        )
    return client

def count_tokens(text, model="gpt-4o"):
    """Count the number of tokens in the text for the specified model."""
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)


def get_pdfs_from_folder(folder_path):
    """Returns a list of PDF file paths from the given folder."""
    if folder_path.lower().endswith('.pdf'):
        return [folder_path]
    
    pdf_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]
    return pdf_files


def extract_docx_content(file_path):
    """Extract text content from a Word document."""
    doc_content = ""
    doc = docx.Document(file_path)
    
    # Extract paragraphs
    for para in doc.paragraphs:
        doc_content += para.text + "\n"

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_content += cell.text + " | "
            doc_content += "\n"
            
    return doc_content


def get_document_name(file_path):
    """Extract the base document name from a file path without extension."""
    base_name = os.path.basename(file_path)
    document_name = os.path.splitext(base_name)[0]
    return document_name


def save_text_to_file(text, file_path):
    """Save text content to a file with proper encoding."""
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(text)
    return file_path

def is_scanned_pdf_from_text(extracted_text, cid_threshold=5):
    scanned_pages = 0
    pages = extracted_text.split("\n\nPage ")

    for page in pages:
        page_content = page.strip()
        if page_content.lower().startswith("page "):
            page_content = page_content[6:].strip()

        if len(page_content) < 200:
            scanned_pages += 1

    cid_count = len(re.findall(r'\(cid:\d+\)', extracted_text))

    return scanned_pages > 2 or cid_count >= cid_threshold

def upload_to_s3(file_path, s3_key, bucket_name):
    """Upload a file to an S3 bucket"""
    s3 = boto3.client(
        's3',
        aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY')
    )
    
    try:
        s3.upload_file(file_path, bucket_name, s3_key)
        print(f"Successfully uploaded {file_path} to s3://{bucket_name}/{s3_key}")
        return True
    except ClientError as e:
        print(f"Error uploading to S3: {str(e)}")
        return False

def upload_contract_to_s3(contract_path, contract_id, excel_file_path=None, json_file_path=None, validation_folder_path=None):
    """
    Uploads contract documents and outputs to S3 in organized structure
    
    Args:
        contract_path (str): Path to contract file or folder
        contract_id (str): Unique identifier for the contract
        excel_file_path (str, optional): Path to generated Excel file
        json_file_path (str, optional): Path to generated JSON file
        validation_folder_path (str, optional): Path to AI validation folder
    
    Returns:
        bool: True if all uploads succeeded, False otherwise
    """
    try:
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        base_folder = os.getenv('AWS_FOLDER_NAME')
        s3_base_path = f"{base_folder}/{contract_id}"
        upload_success = True

        # Upload original documents
        if os.path.isdir(contract_path):
            pdf_documents = get_pdfs_from_folder(contract_path)
            for pdf in pdf_documents:
                original_s3_key = f"{s3_base_path}/original_document/{os.path.basename(pdf)}"
                if not upload_to_s3(pdf, original_s3_key, bucket_name):
                    upload_success = False
        else:
            original_s3_key = f"{s3_base_path}/original_document/{os.path.basename(contract_path)}"
            if not upload_to_s3(contract_path, original_s3_key, bucket_name):
                upload_success = False

        # Upload Excel output if provided
        if excel_file_path and os.path.exists(excel_file_path):
            excel_s3_key = f"{s3_base_path}/output_result/{os.path.basename(excel_file_path)}"
            if not upload_to_s3(excel_file_path, excel_s3_key, bucket_name):
                upload_success = False

        # Upload JSON file if provided
        if json_file_path and os.path.exists(json_file_path):
            json_s3_key = f"{s3_base_path}/output_result/{os.path.basename(json_file_path)}"
            if not upload_to_s3(json_file_path, json_s3_key, bucket_name):
                upload_success = False

        # Upload AI validation files if provided
        if validation_folder_path and os.path.exists(validation_folder_path):
            for root, _, files in os.walk(validation_folder_path):
                for file in files:
                    local_path = os.path.join(root, file)
                    relative_path = os.path.relpath(local_path, validation_folder_path)
                    s3_key = f"{s3_base_path}/Ai_validation/{relative_path}"
                    if not upload_to_s3(local_path, s3_key, bucket_name):
                        upload_success = False

        return upload_success

    except Exception as e:
        print(f"Error in S3 upload process: {str(e)}")
        return False
