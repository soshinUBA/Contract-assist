from openai import OpenAI
import pdfplumber
import docx
from dotenv import load_dotenv
import os
import pandas as pd
import unicodedata
import regex as re
from llama_parse import LlamaParse

load_dotenv(".env")
os.environ["LLAMA_CLOUD_API_KEY"] = os.getenv("LAMA_CLOUD_API_KEY")

pdf_documents = ["./Contracts/test/COMPLETED Cricut, Inc  Other 04 06 2022.pdf", "./Contracts/test/COMPLETED Cricut, Inc  Other 04 06 2022 (1).pdf"]
word_doc = "./FAQ_document_updated.docx"
all_text = ""
faq_doc = ""
parser = LlamaParse(result_type="markdown")

# Load the .docx file
doc = docx.Document(word_doc)
for para in doc.paragraphs:
    faq_doc += para.text + "\n"

# Extract and store tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            faq_doc += cell.text + " | "
        faq_doc += "\n"

print(faq_doc)
with open("faq.txt", "w") as f:
    f.write(faq_doc)
 
print("############# Starting PDF  Processing ######################")

# for pdf_document in pdf_documents:
#     document = parser.load_data(pdf_document)  # Process PDF with LlamaParse
#     all_text += "\n\n".join([doc.text for doc in document])



# Open the PDF file
for pdf_document in pdf_documents:
    with pdfplumber.open(pdf_document) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines = text.split("\n")
                for line in lines:
                    line = line.strip()

                    # Convert potential headers (short uppercase lines)
                    if len(line) < 50 and line.isupper():
                        all_text += f"## {line}\n\n"
                    elif line.startswith("- ") or line.startswith("* "):
                        all_text += f"{line}\n"
                    elif ":" in line:  # Bold key-value pairs
                        key, value = line.split(":", 1)
                        all_text += f"**{key.strip()}**: {value.strip()}\n\n"
                    else:
                        all_text += f"{line}  \n"  # Ensuring Markdown line breaks
            
            # Extract tables (if any) from the page
            tables = page.extract_tables()
            for table in tables:
                if table and len(table[0]) > 0:
                    # Fix: Handle None values in header
                    header = " | ".join([str(cell) if cell is not None else "" for cell in table[0]])
                    separator = " | ".join(["---"] * len(table[0]))

                    all_text += f"\n{header}\n{separator}\n"

                    for row in table[1:]:  # Skip header row
                        # Fix: Handle None values in table rows
                        formatted_row = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        all_text += f"{formatted_row}\n"

with open("contract_b4.md", "w", encoding="utf-8") as f:
    f.write(all_text)
    print(all_text)
def clean_excess(text):
    """
    This function removes any extra lines or names following the extracted customer name.
    It assumes the customer name consists of two or three words, and stops after that.
    """
    lines = text.split('\n')
    # Look for lines that likely contain names, we assume most customer names won't exceed 3 words
    cleaned_lines = []
    for line in lines:
        words = line.split()
        if len(words) > 3:  # If more than 3 words, it likely contains extra info
            cleaned_lines.append(' '.join(words[:3]))  # Take only the first 3 words
        else:
            cleaned_lines.append(line.strip())  # Otherwise, keep the full line if it's valid
    return ' '.join(cleaned_lines).strip()  # Return the cleaned customer name

p_email_pattern = re.compile(r'Email\s*Address:\s*(?:[:]*\s*)?([^\s]+@[^\s]+)')

# Find the primary licensed user email section and extract the email
match = p_email_pattern.search(all_text)
if match:
    primary_user_email = match.group(1).strip()
    print(f"Primary Licensed User Email: {primary_user_email}")
else:
    primary_user_email = "Not found on the document"
    print("Primary Licensed User Email not found.")


def replace_partial_name(all_text, customer_name, new_name):
    # Split the customer name into words
    customer_name_words = customer_name.split()

    # Only proceed if the customer name has 3 or more words
    if len(customer_name_words) >= 3:
        first_two_words = ' '.join(customer_name_words[:2])  # Get the first two words
        last_two_words = ' '.join(customer_name_words[-2:])  # Get the last two words

        # Compile regex patterns for the first two and last two words to handle possessive forms
        first_two_pattern = re.compile(re.escape(first_two_words) + r"(?:’s|'s)?", re.IGNORECASE)
        last_two_pattern = re.compile(re.escape(last_two_words) + r"(?:’s|'s)?", re.IGNORECASE)

        # Replace the first two and last two words with the new name
        all_text = re.sub(first_two_pattern, new_name, all_text)
        all_text = re.sub(last_two_pattern, new_name, all_text)

    return all_text

run_open_ai = True
# First regex pattern to extract the customer name from the first section
customer_name_pattern = r"Customer Billing / Third Party Payor Monotype\s*\(if same as Customer, indicate below\)\s*Name:\s*([\w\s]+?)(?:\s*Same as Customer)?\s*(?=\s*Contact:|\s*Name:|\n)"
customer_name_pattern_2 = r"Name:\s*([A-Za-z\s]+)\s*Same as Customer"
customer_name_pattern_3 = r"Name:\s*([A-Za-z\s]+)\s*\|\s*\|\s*\|\s*Name:\s*([A-Za-z\s]+)"
customer_name_pattern_4 = r"Name:\s*([A-Za-z\s\.]+)\s*Name:\s*Same as Customer"

# Extract the customer name using the first regex pattern
match = re.search(customer_name_pattern, all_text, re.DOTALL)
match2 = re.search(customer_name_pattern_3, all_text, re.DOTALL)
match3 = re.search(customer_name_pattern_4, all_text, re.DOTALL)
if match:
    customer_name = match.group(1).strip()  # Extract the customer name, remove leading/trailing whitespace
    print(f"Customer name found: {customer_name}")
    print("1st")
    # Step 1: Normalize the text to handle line breaks (optional, for safe replacement)
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Step 2: Replace the found customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

elif match3:
    customer_name = match3.group(1).strip()  # Extract and clean the customer name
    print(f"Customer name found: {customer_name}")
    print("2nd")

    # Step 1: Normalize the text to handle line breaks
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Step 2: Replace the found customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

# New if statement to match the second pattern
elif re.search(customer_name_pattern_2, all_text, re.DOTALL):
    match = re.search(customer_name_pattern_2, all_text, re.DOTALL)
    customer_name = match.group(1).strip()  # Extract and clean the customer name
    print(f"Customer name found: {customer_name}")
    print("3rd")

    # Normalize text to handle line breaks
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Replace the customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

# Elif for Agreement and Monotype Imaging Inc. extraction
elif re.search(r"Agreement are made part of this Agreement\.\s*Monotype Imaging Inc\.\s*([A-Za-z\s]+?)(?=\s*\n[A-Z])", all_text, re.DOTALL):
    match = re.search(r"Agreement are made part of this Agreement\.\s*Monotype Imaging Inc\.\s*([A-Za-z\s]+?)(?=\s*\n[A-Z])", all_text, re.DOTALL)
    customer_name = match.group(1).strip()  # Extract the customer name
    print(f"Customer name found before cleaning: {customer_name}")
    print("4th")

    # Post-process the customer name to remove excess information
    customer_name_cleaned = clean_excess(customer_name)
    print(f"Customer name found after cleaning: {customer_name_cleaned}")

    # Step 1: Normalize the text to handle line breaks (optional, for safe replacement)
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Step 2: Replace the found customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name_cleaned, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

elif match2:
    customer_name = match2.group(1).strip()  # Extract the customer name from the first "Name:" entry
    print(f"Customer name found: {customer_name}")
    print("5th")

    # Normalize the text (optional)
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Replace the found customer name with a new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

elif re.search(r"Name:\s*([A-Za-z\s]+)\s*\|\s*\|.*?Name:\s*([A-Za-z\s]+)", all_text, re.DOTALL):
    match = re.search(r"Name:\s*([A-Za-z\s]+)\s*\|\s*\|.*?Name:\s*([A-Za-z\s]+)", all_text, re.DOTALL)
    customer_name = match.group(1).strip()  # Extract the first customer name
    print(f"Customer name found: {customer_name}")
    print("6th")

    # Step 1: Normalize the text to handle line breaks
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Step 2: Replace the found customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

elif re.search(r"Name:\s*([A-Za-z\s\.]+)\s*(?=Name:)", all_text, re.DOTALL):
    match = re.search(r"Name:\s*([A-Za-z\s\.]+)\s*(?=Name:)", all_text, re.DOTALL)
    customer_name = match.group(1).strip()  # Extract the customer name
    print(f"Customer name found: {customer_name}")
    print("7th")

    # Normalize text to handle line breaks
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    # Replace the found customer name with the new name
    new_name = "Blessing Company Ltd"
    all_text_normalized = all_text_normalized.replace(customer_name, new_name)

    # Call the function to replace first two and last two words
    all_text_normalized = replace_partial_name(all_text_normalized, customer_name, new_name)

    # Update the original text with the normalized and replaced content
    all_text = all_text_normalized

else:
    print("Customer name not found.")
    customer_name = "Delta Innovations Triple"
    new_name = "Blessing Company Ltd"

    # Normalize line breaks (fixing words split by newlines)
    all_text_normalized = re.sub(r'([a-zA-Z])\s*\n\s*([a-zA-Z])', r'\1 \2', all_text)

    all_text_normalized = all_text_normalized.replace(customer_name, new_name)
    # run_open_ai = False

with open("contractb4Email.txt", "w", encoding="utf-8") as f:
    f.write(all_text)

# Function to normalize Unicode text to ensure consistency
def normalize_text(text):
    return unicodedata.normalize('NFKC', text)

# Normalize all_text before processing to handle special characters
all_text = normalize_text(all_text)

# Updated email pattern to handle Unicode characters in email addresses
email_pattern = r'([\w._%+-]+(?:\s*\n*\s*)@(?:\s*\n*\s*)[\w.-]+\.[a-zA-Z]{2,})'

# Alternatively, use Unicode property for more complex characters like 'ö'
unicode_email_pattern = r'([\w\p{L}._%+-]+(?:\s*\n*\s*)@(?:\s*\n*\s*)[\w.-]+\.[a-zA-Z]{2,})'

# Use the above pattern to find all emails in the text
emails = re.findall(unicode_email_pattern, all_text)

# Debugging to check emails found after updating the pattern
print(f"Updated Found emails: {emails}")

# Now process and replace the emails similarly to how you did before
for email in emails:
    normalized_email = re.sub(r'\s+', '', email)  # Normalize by removing spaces/newlines
    if not normalized_email.endswith('@monotype.com'):
        print(f"Replacing customer email: {normalized_email}")
        new_email = "new.email@example.com"
        all_text = re.sub(re.escape(normalized_email), new_email, all_text)

# After the initial pass, find and handle any remaining emails that may still need domain updates
emails2 = re.findall(email_pattern, all_text)

for email in emails2:
    normalized_email = re.sub(r'\s+', '', email)  # Normalize again just in case
    if not (normalized_email.endswith("@monotype.com") or normalized_email.endswith("@example.com")):
        # Extract the domain part from the customer email
        domain_pattern = r'@([a-zA-Z0-9.-]+)\.[a-zA-Z]{2,}'
        customer_domain = re.search(domain_pattern, normalized_email)
        if customer_domain:
            domain_to_replace = customer_domain.group(1)  # e.g., 'brightstarpublishing'
            # Construct a regex that matches the domain part only inside the email
            domain_with_email_pattern = r'(?<=@)' + re.escape(domain_to_replace) + r'(?=\.[a-zA-Z]{2,})'
            # Replace the domain part only inside the email
            all_text = re.sub(domain_with_email_pattern, 'example', all_text)

# Find broken or special character emails and reassemble them
broken_email_pattern = r'([a-zA-Z0-9._%+-]+)\s*\n*\s*(@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
# Reassemble broken emails
all_text = re.sub(broken_email_pattern, r'\1\2', all_text)

# Replace reconstructed emails like 'egallagher@example.com' with the desired fixed email
all_text = re.sub(r'[a-zA-Z0-9._%+-]+@example\.com', 'new.email@example.com', all_text)

name_patterns = [
    # 1. Primary Licensed Monotype Fonts User Section, allowing comma after the name and Unicode letters
    r"Primary Licensed Monotype Fonts User\s*\|\s*Name:\s*([\p{L}]+\s+[\p{L}]+),?\s",

    # 2. Prepared by: Contact Section
    r"Contact:\s*([\p{L}]+\s+[\p{L}]+)\s*\|\s*\|",

    #2.5 for contact
    r"Contact:\s*([\p{L}]+\s+[\p{L}]+)(?=\s*Contact|E-mail)",

    # 3. Name in Monotype Limited or Company section
    r"\nMonotype Limited.*?(?:\n|)\s*([\p{L}]+\s+[\p{L}]+)\s*([\p{L}]+\s+[\p{L}]+)"
]

# Function to extract and replace all occurrences of names
def replace_names(text, name_patterns):
    # Find and store all names to replace
    all_names = []
    for pattern in name_patterns:
        matches = re.findall(pattern, text)
        if matches:
            for match in matches:
                # Ensure we're capturing names from multiple groups if applicable
                if isinstance(match, tuple):
                    all_names.extend(match)
                else:
                    all_names.append(match)

    # Replace each found name with "John Doe" throughout the entire text
    for name in set(all_names):  # Use `set` to avoid duplicate replacements
        print(f"Replacing name: {name}")
        text = re.sub(re.escape(name), "John Doe", text)
    
    return text

# Apply the function
all_text = replace_names(all_text, name_patterns)


with open("contract.txt", "w", encoding="utf-8") as f:
    f.write(all_text)

print("############# Starting Openai Processing ######################")


if run_open_ai:
    client = OpenAI(
    api_key=os.getenv("openai-api-key")  # this is also the default, it can be omitted
    )

    completion = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        max_tokens=3500,
        messages=[
            {"role": "system", "content": """
                YOU WILL BE HEAVILY PENALIZED FOR NOT FOLLOWING INSTRUCTIONS. 
                - Upon receiving a user inquiry, you **must** directly search the document and return answers for all 24 fields in the required format. 
                - **Do not** engage in a back-and-forth conversation, explain what you are doing, or restate the user's query.
                - Your response **must only be the final answer** in tabular form, with the 24 field names and corresponding values. Do **not** return intermediate questions, commentary, or processing steps.
                - If a field is not found, respond with 'Not found on this document' for that specific field. 
                - Do **not** start by saying 'Let's search for this first' or any variation. Simply return the completed table at the end.
                - You are **not allowed** to ask questions. Only perform the task as instructed and return the output in one go.
                Strict rules to follow:
                    1. Ask no questions or restate no queries. Your task is to **silently** process the document and deliver the final output.
                    2. Return only the final table with answers for all 24 fields. Intermediate steps or explanations will lead to penalties.
                    3. Format your response in this table structure: 'Field, Value, Reason for value' (where applicable).
                    4. If the information is not found in the document, write 'Not found on this document' for that field.
                    5. Always adhere strictly to these instructions and focus solely on delivering the expected tabular output.

                "Monotype Fonts Pro" IS NOT "Monotype Fonts Support" DO NOT CONFUSE THEM.
                "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Preferred Service​"
                "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "Monotype Fonts Pro"
                "YOU WILL BE HEAVILY PENALIZED IF YOU RETURN THIS RESULT "All Font Software available on Monotype Fonts during the Term."
                Ensure that you:
                Search the specific customer’s PDF for each field’s answer. Do not refer to other documents or sources.
                Do not send any results until all 24 questions have been answered. Ensure accuracy is maintained for each answer, and if any information is not found in the document, respond with "Not found on this document."
                Do not provide partial or incomplete answers. Ensure both "Name Fonts" and "Material Number" are fully answered before presenting the response. DO NOT GIVE random sentences as answers or incomplete values or "refer to the documents" or "and more" or "etc" type of values
                Note: customer_name is a placeholder for the actual customer name provided in the user’s query. Ensure all answers are specific to the customer referenced in the query.
                Important Addendum Consideration:
                In your investigation, always review any addendums related to the contract. Addendums modify specific terms in the original contract and supersede any corresponding terms in the original document. If an addendum alters a term or clause, the addendum takes precedence. For any terms not mentioned in the addendum, the original contract remains in effect. Ensure you cross-check the contract and addendum carefully to provide the most accurate and up-to-date information.
                Process:
                For each of the following questions, search the entire customer’s PDF document for the field seperately else you will be heavily penalized.
                Do not guess or assume answers. If the answer is not available, respond with "Not found on this document." 
                #####
                Questions for Each Field:
                1.	Contract Type: "What is the contract type for customer_name (something something Order Form, The full value would be along the lines of 'Monotype Fonts Service and License Order Form')?"
                2.	Contract End Date: "What is the contract end date for customer_name?"
                3.	Contract Start Date: "What is the contract start date for customer_name?"
                4.	Contract Number: "What is the unique contract number for customer_name?"
                5.	Customer Name: "What is the entity name agreeing to the contract?"
                6.	Web Page Views: "What is the Licensed Page Views (Web Page Content) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
                7.	Digital Ad Views: "What is the Licensed Impressions (Digital Marketing Communications) for customer_name? Search your knowledge (This can be found in the License Usage per Term section)"
                8.	Commercial Documents: "What is the upper limit for licensed commercial documents for customer_name?"
                9.	Licensed Applications: "How many Licensed Applications does customer have (DO NOT FETCH Software Products)? (Do not confuse it with other licensed components like Licensed Software Products, or licensed desktop application else you will be heavily penalized) "
                10.	Licensed User Count: "What is the total number of Monotype fonts portal users for customer_name?"
                11.	Monotype Font Support: "{Which Monotype fonts support level did 'customer_name' choose: basic, premier, elite or "Not found on the document" ? DON'T GIVE ANY ANSWER APART FORM THESE 3. If "Monotype Font Support" is not explicitly mentioned the answer should be "Not found on the document" else it should be the value under "Monotype Fonts Support". DO NOT RETURN "Yes" or "No" for this field}"
                12.	Primary Licensed User: "Who is the primary licensed user and what is their name and title?"
                13.	Add-On Type:  Always answer 'Not found on this document' 
                14.	Name Fonts: "List all the font names for 'customer_name'. Do not include any incomplete values, vague statements or sentances such as 'refer to the documents', 'and more', 'All Font Software available on Monotype Fonts during the Term.', etc else you be very heavily penalized. Return only the exact font names."
                15.	Material Number: Return this only if the "Name Fonts" field has an output. "What are the material numbers for customer_name?. GIVE ALL THE VALUES, DO NOT GIVE PARTIAL VALUES" 
                16.	Contract Name: "What is the name of the contract document for customer_name?"
                17.	Offline Contract: Default is "Yes", unless specified "Online". "Is this an offline contract for customer_name? (Store 'Yes' if no indication of being online)."
                18.	Opportunity ID: "Is there an Opportunity ID for customer_name? (Store 'No' if unavailable)."
                19.	SAP ID: "Is there an SAP ID for customer_name? (Store 'Not found on this document' if unavailable)."
                20.	Territory: "What is the country listed in the address for customer_name? (The country should not be abbreviated and is found ONLY IN THE CUSTOMER ADDRESS. DO NOT TAKE ANY OTHER ADDRESS OTHER THAN THE CUSTOMER)."
                21.	Additional User Count: "How many Additional Licensed Desktop Users (which are not Licensed Monotype Fonts Users) can the customer_name have? DO NOT COUFUSE THIS WITH "Licensed Desktop Users" THEY ARE NOT THE SAME"
                22.	Swapping Allowed: "{"Can production fonts be swapped for 'customer_name'? Answer 'Yes' or 'No' strictly based on the Production Fonts field in the License for Monotype Fonts License Terms. Only answer 'Yes' if the terms 'swap' or 'replace' are explicitly mentioned in relation to Production Fonts. If not, the answer should be 'No.' Do not infer the answer."}"
                23: Production font: "How many production fonts does customer have in contract as well as addendum, if present?" Give the exact words/paragraph as in the contract else you will be heavily penalized. This is only found in the "License Usage per Term" section.
                24. Reporting days: "How many days does customer_name have to report their usage of the font software as Production Fonts? Focus only on the reporting days explicitly mentioned for reporting Production Fonts usage after receiving the list of downloaded Font Software. Exclude any references to providing information upon request or disputing inaccuracies. Include any additional days granted only if they follow a formal notice. Provide the result in the format 'X days and Y additional days' if applicable, or just 'X days' if no additional days are mentioned"
                25. Licensed Electronic Documents: "What is the number of Commercial Electronic Documents the customer_name can have?"
                26. Licensed Externally Accessed Servers: "How many Externally Accessed Servers can the customer have?"
                27. Binding Obligations: "What is the Binding Obligation that is bound to this contract?"
                28. Primary User Email: "What is the email address of the primary licensed user for customer_name?"
                29. Primary User First Name: "What is the first name of the primary user?"
                30. Primary User Last Name: "What is the last name of the primary user?"
                31. Brand & License Protection: "What is the Brand & License Protection for the customer?"
                32. Past Usage Term(Dates): "What is the past usage terms from the customer (only give dates, if it exist)?"  
                
             
                #####
                IMPORTANT POINTS: 
                * LICENSED APPLICATION AND LICENSED SOFTWARE PRODUCT ARE DIFFERENT.
                1. PROCESS for each answer Separately else you will be very heavily penalized.
                2. Important point: "Always refer to Documents before answering Questions"
                3. USE MEDIUM CHUNKING SIZE else you will be heavily penalized
                4. YOU ARE NOT ALLOWED TO REFER TO ANY EXTERNAL SOURCES
                5. Important Point: "PRODUCTION FONTS, PRODUCTION SOFTWARE ARE SAME AND CAN BE REFERRED VICE VERSA."
                7. YOU CANNOT CHANGE THE INFORMATION, ELSE YOU WILL BE HEAVILY PENAILZED.
                8. WHILE FETCHING MULTIPLE CONTRACTS DATA, YOU SHOULD UNDERSTAND AND FETCH DATA FROM ALL CONTRACTS INDIVISUALLY.
                9. Signing and preparation of contract are different things. It should be treated differently.
                10. ALWAYS SEARCH YOUR KNOWLEDGE BEFORE ANSWERING THE QUESTION
                11. Always Answer in SHORT EXCEPT FOR THE FOLLOWING FIELDS, Name Fonts, Material Number, Production Font and Reporting. For these 4 fields your answer can be LONG
                12. The Contract's Date whose termination date is before today's date is expired.
                13. YOU CANNOT MAKE ANY ASSUMPTIONS.
                15. ALWAYS present fields, their values and their reasons in a clear tabular format. Ensure that all answers are contained within the table and not outside of it.
                16. Your output should EXACTLY BE THE 24 fields MENTIONED or you will be penalized. A sample output for the 24 fields from the FAQ document. Ensure that all answers are specific to the correct customer and contract in question.
                18. YOU SHOULD FOLLOW ALL ABOVE MENTIONED POINTS ELSE YOU WILL BE PENAILZED HEAVILY
                sample output format, Field,Value,Reason for value
            """},
            {
                "role": "user",
                "content": f""" 
                    This is the faq document: {faq_doc}.
                    ######
                    Give me the details of the contract below:

                    {all_text} 
                """
            }
        ]
    )

    message = completion.choices[0].message.content
    pdf_document = pdf_documents[0]

    # Extract the base name of the file (e.g., M00217189.pdf)
    base_name = os.path.basename(pdf_document)

    # Remove the .pdf extension to get just the document name (e.g., M00217189)
    document_name = os.path.splitext(base_name)[0]

    # Create the desired output file name (e.g., M00217189_Output.txt)
    output_filename = f"./Output/{document_name}_Output.txt"
    with open(output_filename, "w", encoding="utf-8") as file:
        file.write(message)  # Assuming 'message' contains the content you want to write

    print("Message content saved successfully!")
    print(completion.choices[0].message.content)

    # Step 1: Parse the content into a structured format (table)
    lines = [line for line in completion.choices[0].message.content.splitlines() if '----' not in line]  # Remove separator lines

    fields, values, reasons = [], [], []

    for line in lines[1:]:  # Start from index 1 to skip the column headers
        if '|' in line:  # Ensure it's a valid row
            parts = line.split('|')[1:4]  # Split into Field, Value, and Reason
            if len(parts) == 3:  # Ensure there are exactly three parts
                fields.append(parts[0].strip())
                values.append(parts[1].strip())
                reasons.append(parts[2].strip())

    # Ensure email logic works correctly
    if len(values) > 27 and ("new.email@example.com" in values[27] or "@example.com" in values[27] or "@example" in values[27]):
        values[27] = primary_user_email
        print("Changed the primary user email")
    else:
        print("Failed to change primary user email")

    # Update customer name
    if len(values) > 3:
        values[4] = customer_name

    # Step 2: Create a DataFrame
    data = {
        "Field": fields,
        "Value": values,
        "Reason for value": reasons
    }
    df = pd.DataFrame(data)

    # Step 3: Export to Excel
    file_path = f'./ExcelOutput/{document_name}_data.xlsx'
    df.to_excel(file_path, index=False)

    print(f"Data has been exported to {file_path}")